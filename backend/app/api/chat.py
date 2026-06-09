"""
对话接口路由
提供消息发送、SSE 流式对话、历史记录查询等功能
"""

import json
import logging
import os
import time
import asyncio
import re
from pathlib import Path

import httpx
from fastapi import APIRouter, Depends, HTTPException, Query, status, UploadFile
from fastapi.responses import FileResponse, PlainTextResponse
from sqlalchemy.orm import Session
from sse_starlette.sse import EventSourceResponse

from app.api.auth import get_current_user
from app.config import settings
from app.database import get_db
from app.models.conversation import Conversation
from app.models.message import Message
from app.models.user import User
from app.schemas.chat import (
    ConversationCreateRequest,
    ConversationDetailResponse,
    ConversationListResponse,
    ConversationResponse,
    MessageCreateRequest,
    MessageResponse,
)
from app.schemas.common import APIResponse

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/chat", tags=["对话"])

ARTIFACT_ROOT = Path("artifacts")


def _artifact_dir(user_id: int, conversation_id: int) -> Path:
    path = ARTIFACT_ROOT / str(user_id) / str(conversation_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _safe_artifact_name(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", name).strip(" .")
    return cleaned[:120] or "artifact"


def _write_text_artifact(user_id: int, conversation_id: int, name: str, content: str, suffix: str) -> Path:
    directory = _artifact_dir(user_id, conversation_id)
    stem = _safe_artifact_name(name)
    path = directory / f"{stem}{suffix}"
    counter = 2
    while path.exists():
        path = directory / f"{stem}_{counter}{suffix}"
        counter += 1
    path.write_text(content or "", encoding="utf-8")
    return path


def _write_report_artifact(user_id: int, conversation_id: int, name: str, markdown: str) -> Path:
    """Write a report as docx when python-docx is available, otherwise markdown."""
    directory = _artifact_dir(user_id, conversation_id)
    stem = _safe_artifact_name(name)
    try:
        from docx import Document

        path = directory / f"{stem}.docx"
        counter = 2
        while path.exists():
            path = directory / f"{stem}_{counter}.docx"
            counter += 1
        doc = Document()
        doc.add_heading(name, level=1)
        for line in (markdown or "").splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s+", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            else:
                doc.add_paragraph(stripped)
        doc.save(path)
        return path
    except Exception as exc:
        logger.warning("写入 docx 产物失败，降级为 md: %s", exc)
        return _write_text_artifact(user_id, conversation_id, name, markdown, ".md")


def _preview_artifact_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document

            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return "该 Word 文档暂时无法预览，可下载后查看。"
    return path.read_text(encoding="utf-8", errors="ignore")


def _derive_conversation_title(content: str) -> str:
    text = re.sub(r"[#>*_`~\[\]（）(){}|]+", " ", content or "")
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return "文件分析"
    return text[:24] + ("..." if len(text) > 24 else "")


def _file_url_to_local_path(file_url: str) -> str:
    local_path = file_url.lstrip("/")
    if not os.path.isabs(local_path):
        local_path = os.path.join(os.getcwd(), local_path)
    return local_path


def _is_audio_file(file_url: str | None) -> bool:
    if not file_url:
        return False
    return os.path.splitext(file_url)[1].lower() in {
        ".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac", ".mp4", ".avi", ".mov", ".mkv", ".webm"
    }


def _is_resume_file(file_url: str | None) -> bool:
    if not file_url:
        return False
    return os.path.splitext(file_url)[1].lower() in {
        ".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png", ".webp", ".bmp"
    }


def _parse_file_urls(file_url: str | None, file_urls: str | None) -> list[str]:
    urls: list[str] = []
    if file_urls:
        for item in file_urls.split(","):
            cleaned = item.strip()
            if cleaned and cleaned not in urls:
                urls.append(cleaned)
    if file_url and file_url not in urls:
        urls.insert(0, file_url)
    return urls


def _parse_file_names(file_names: str | None, file_urls: list[str]) -> list[str]:
    names: list[str] = []
    if file_names:
        for item in file_names.split(","):
            cleaned = item.strip()
            if cleaned:
                names.append(cleaned)
    if len(names) < len(file_urls):
        names.extend(Path(url).name for url in file_urls[len(names):])
    return names[:len(file_urls)]


def _select_file_for_agent(task: dict, file_urls: list[str]) -> str | None:
    assigned = task.get("file_urls")
    if isinstance(assigned, list) and assigned:
        return str(assigned[0])

    agent_type = task.get("agent_type")
    if agent_type == "resume":
        return next((url for url in file_urls if _is_resume_file(url)), None)
    if agent_type == "recording":
        return next((url for url in file_urls if _is_audio_file(url)), None)
    if agent_type == "question":
        return (
            next((url for url in file_urls if _is_resume_file(url)), None)
            or next((url for url in file_urls if _is_audio_file(url)), None)
        )
    return None


async def _stream_content_chars(
    content: str,
    agent_name: str | None = None,
    metadata: dict | None = None,
):
    """Emit answer text as small SSE chunks for a typing effect."""
    text = content or ""
    for index, char in enumerate(text):
        payload = {
            "type": "content",
            "content": char,
            "done": False,
        }
        if agent_name:
            payload["agent_name"] = agent_name
        if metadata and index == len(text) - 1:
            payload["metadata"] = metadata
        yield json.dumps(payload, ensure_ascii=False)
        await asyncio.sleep(0.01)


async def _web_search(query: str, max_results: int = 5) -> list[dict[str, str]]:
    """Use Tavily for web references when configured."""
    if not settings.TAVILY_API_KEY or not query.strip():
        return []

    try:
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": settings.TAVILY_API_KEY,
                    "query": query,
                    "max_results": max_results,
                    "include_answer": True,
                    "search_depth": "advanced",
                },
            )
        if resp.status_code != 200:
            return []
        return [
            {
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": item.get("content", ""),
            }
            for item in resp.json().get("results", [])
        ]
    except Exception as e:
        logger.warning(f"联网搜索失败: {e}")
        return []

# 面试助手的系统提示词
SYSTEM_PROMPT = """你是"智面"——一个温暖、专业、充满鼓励的AI面试助手。你的使命是陪伴每一位求职者走过面试准备的每一步。

## 你的性格
- 热情真诚，像一位经验丰富又贴心的学长/学姐
- 善于鼓励，在给出建议的同时肯定用户的努力和进步
- 回答具体实用，绝不给空泛的套话，每一条建议都要能落地执行
- 适时使用轻松的语气，让紧张的面试准备变得有趣

## 你的核心能力
1. **简历优化**：逐条分析简历内容，指出亮点和不足，给出具体的改进措辞建议
2. **面试模拟**：扮演面试官进行模拟面试，提问后给出参考回答和评分
3. **面试题精讲**：不只是给答案，而是讲清楚"为什么这么答"、"面试官想考察什么"
4. **行业洞察**：分享目标公司和岗位的真实面试风格、高频考点
5. **心理支持**：面试前帮用户缓解紧张，建立自信

## 回答规范
- 使用 Markdown 格式让回答结构清晰（标题、列表、加粗重点）
- 开头用简短的话回应或鼓励用户，再进入正式内容
- 举例时尽量贴合用户的实际背景（从对话历史中了解）
- 结尾给一个温暖的鼓励或下一步行动建议
- 适当使用表情符号增加亲和力，但不要过度"""


async def _build_user_context(db: Session, user_id: int) -> str:
    """
    构建用户个人信息上下文，用于注入 system prompt
    
    Args:
        db: 数据库会话
        user_id: 用户ID
        
    Returns:
        str: 格式化后的用户个人信息文本，若无需注入则返回空字符串
    """
    from app.models.user_profile import UserProfile

    profile = db.query(UserProfile).filter(
        UserProfile.user_id == user_id
    ).first()

    if not profile or not profile.allow_ai_use:
        return ""

    parts = []
    gender_map = {"male": "男", "female": "女", "other": "其他"}

    if profile.age:
        parts.append(f"- 年龄：{profile.age}岁")
    if profile.gender:
        parts.append(f"- 性别：{gender_map.get(profile.gender, profile.gender)}")
    if profile.university:
        parts.append(f"- 学校：{profile.university}")
    if profile.major:
        parts.append(f"- 专业：{profile.major}")
    if profile.grade:
        parts.append(f"- 年级：{profile.grade}")
    extra_labels = {
        "education_level": "学历",
        "degree": "学位",
        "graduation_year": "毕业年份",
        "target_position": "目标岗位",
        "target_city": "目标城市",
        "skills": "技能关键词",
        "certificates": "证书/资格",
        "internship_experience": "实习/项目经历",
    }
    for field, label in extra_labels.items():
        value = getattr(profile, field, None)
        if value:
            parts.append(f"- {label}：{value}")

    if not parts:
        return ""

    return (
        "\n\n## 当前用户个人信息\n"
        + "\n".join(parts)
        + "\n\n在回答时请结合用户的以上背景信息，给出更有针对性的建议。"
        + "\n如果用户询问的问题与以上信息无关，可以忽略。"
    )


async def _generate_resume_sse(
    conversation_id: int,
    file_url: str,
    user_message: str,
    user_id: int,
    db: Session,
) -> str:
    """
    简历评估 SSE 流式生成器
    调用 Qwen-VL 模型评估简历
    """
    from app.agents.resume_agent import ResumeAgent
    import os

    agent = ResumeAgent()

    # 将 URL 转为本地文件路径
    # file_url 格式: /uploads/{user_id}/{filename}
    local_path = file_url.lstrip('/')
    if not os.path.isabs(local_path):
        local_path = os.path.join(os.getcwd(), local_path)

    logger.info(f"简历评估: file_url={file_url}, local_path={local_path}, exists={os.path.exists(local_path)}")

    try:
        # 获取用户个人信息上下文
        user_context = await _build_user_context(db, user_id)
        report_parts: list[str] = []

        async for chunk in agent.evaluate_stream(local_path, user_message or None, user_context or ""):
            if chunk["type"] == "thinking":
                # Qwen-VL 的 reasoning_content 是内部推理，可能包含未确认的岗位猜测；
                # 前端只展示明确的状态与正式报告，避免把中间推测误当成结论。
                continue
            elif chunk["type"] == "content":
                report_parts.append(chunk["content"])
                async for text_chunk in _stream_content_chars(chunk["content"], agent_name="简历评估"):
                    yield text_chunk
            elif chunk["type"] == "status":
                yield json.dumps({
                    "type": "status",
                    "agent_name": "简历评估",
                    "message": chunk.get("message", "处理中..."),
                    "progress": chunk.get("progress", 0),
                }, ensure_ascii=False)
        report_text = "".join(report_parts).strip()
        if report_text:
            _write_report_artifact(user_id, conversation_id, "简历评估报告", report_text)
    except Exception as e:
        logger.error(f"简历评估失败: {e}", exc_info=True)
        async for text_chunk in _stream_content_chars(f"[简历评估失败: {str(e)}]", agent_name="简历评估"):
            yield text_chunk

    yield json.dumps({"type": "content", "content": "", "done": True}, ensure_ascii=False)


async def _generate_recording_sse(
    conversation_id: int,
    file_url: str,
    user_message: str,
    user_id: int,
    db: Session,
) -> str:
    """
    录音分析 SSE 流式生成器
    调用 Paraformer 转文字 + DeepSeek 分析
    """
    from app.agents.recording_agent import RecordingAgent
    import os

    agent = RecordingAgent()

    # 将 URL 转为本地文件路径
    local_path = file_url.lstrip('/')
    if not os.path.isabs(local_path):
        local_path = os.path.join(os.getcwd(), local_path)

    try:
        # 获取用户个人信息上下文
        user_context = await _build_user_context(db, user_id)

        # 先发送状态：正在转写
        yield json.dumps({
            "type": "status",
            "agent_name": "录音分析",
            "message": "正在转写录音...",
            "progress": 30,
        }, ensure_ascii=False)

        result = await agent.process(local_path, user_context or "", user_id=user_id)

        # 发送状态：转写完成
        yield json.dumps({
            "type": "status",
            "agent_name": "录音分析",
            "message": "转写完成，正在分析...",
            "progress": 70,
        }, ensure_ascii=False)

        # 构建分析报告（Markdown 格式）
        if isinstance(result, dict):
            report_parts = []

            # 转写内容
            if result.get("transcription"):
                _write_text_artifact(user_id, conversation_id, "录音转写内容", result["transcription"], ".txt")
                yield json.dumps({
                    "type": "status",
                    "agent_name": "录音分析",
                    "message": "录音转写已保存到右侧产物栏",
                    "progress": 75,
                }, ensure_ascii=False)

            # 音频信息
            duration = result.get("duration_seconds", 0)
            if duration > 0:
                report_parts.append(
                    f"\n## ℹ️ 音频信息\n\n"
                    f"- 时长：{duration // 60} 分 {duration % 60} 秒\n"
                )

            # 分析结果
            if result.get("analysis"):
                analysis = result["analysis"]
                if isinstance(analysis, dict):
                    # 新格式：analysis 包含 analysis_text 字段（Markdown）
                    if analysis.get("analysis_text"):
                        report_parts.append(f"\n{analysis['analysis_text']}")
                    else:
                        # 兼容旧格式（JSON dict）
                        report_parts.append("\n## 📊 分析结果")
                        for key, value in analysis.items():
                            report_parts.append(f"\n**{key}**: {value}")
                else:
                    report_parts.append(f"\n{analysis}")

            content = "\n".join(report_parts)
            if content.strip():
                _write_report_artifact(user_id, conversation_id, "面试录音评估报告", content)
        else:
            content = str(result)
            if content.strip():
                _write_report_artifact(user_id, conversation_id, "面试录音评估报告", content)

        async for text_chunk in _stream_content_chars(content, agent_name="录音分析"):
            yield text_chunk
    except Exception as e:
        logger.error(f"录音分析失败: {e}", exc_info=True)
        async for text_chunk in _stream_content_chars(f"[录音分析失败: {str(e)}]", agent_name="录音分析"):
            yield text_chunk

    yield json.dumps({"type": "content", "content": "", "done": True}, ensure_ascii=False)


async def _generate_question_sse(
    conversation_id: int,
    user_message: str,
    user_id: int,
    db: Session,
    file_url: str | None = None,
) -> str:
    """面试题生成 SSE 生成器，支持结合简历内容个性化出题。"""
    from app.agents.question_agent import QuestionAgent

    yield json.dumps({
        "type": "status",
        "agent_name": "面试题生成",
        "message": "正在检索题库并生成个性化面试题...",
        "progress": 35,
    }, ensure_ascii=False)

    resume_text = ""
    if file_url:
        try:
            local_path = _file_url_to_local_path(file_url)
            if _is_audio_file(file_url):
                from app.agents.recording_agent import RecordingAgent

                yield json.dumps({
                    "type": "status",
                    "agent_name": "面试题生成",
                    "message": "正在转写录音，用于生成相似面试题...",
                    "progress": 20,
                }, ensure_ascii=False)
                trans_result = await RecordingAgent().transcribe(local_path)
                transcript = trans_result.get("text", "")
                if transcript:
                    resume_text = f"以下是面试录音转写内容，请根据其中的面试场景、问题和回答生成相似面试题：\n{transcript}"
            else:
                from app.agents.resume_agent import ResumeAgent

                resume_text = await ResumeAgent()._extract_text(local_path)
        except Exception as e:
            logger.warning(f"文件内容提取失败，降级为普通出题: {e}")

    agent = QuestionAgent()
    questions = await agent.generate(
        category="general",
        difficulty="medium",
        count=8,
        topic=user_message or "大学生求职面试",
        resume_text=resume_text or None,
    )
    saved_ids = await agent.save_to_db(
        questions,
        category="general",
        difficulty="medium",
        user_id=user_id,
    )

    try:
        from app.services.rag_service import RAGService

        rag = RAGService()
        await rag.add_documents(
            documents=[q.get("question", "") for q in questions],
            metadatas=[
                {"category": q.get("category", "general"), "difficulty": q.get("difficulty", "medium")}
                for q in questions
            ],
        )
    except Exception:
        pass

    lines = ["## 个性化面试练习题\n"]
    for idx, q in enumerate(questions, 1):
        if not isinstance(q, dict):
            continue
        question_text = str(q.get("question", "")).strip()
        if not question_text:
            continue
        lines.append(f"### {idx}. {question_text}")
        key_points = q.get("key_points") or []
        if key_points:
            lines.append("**考察点：**")
            lines.extend(f"- {str(p)}" for p in key_points)
        answer = q.get("reference_answer")
        if answer:
            lines.append("**参考答案：**")
            lines.append(str(answer).strip())
        lines.append("")

    yield json.dumps({
        "type": "status",
        "agent_name": "面试题生成",
        "message": "题目已生成并写入题库",
        "progress": 100,
    }, ensure_ascii=False)
    async for text_chunk in _stream_content_chars(
        "\n".join(lines),
        agent_name="面试题生成",
        metadata={"question_ids": saved_ids},
    ):
        yield text_chunk
    yield json.dumps({"type": "content", "content": "", "done": True}, ensure_ascii=False)


async def _resolve_agent_decision(agent_type: str, content: str, file_urls: list[str]) -> dict:
    """统一入口：由主 Agent 生成调度计划。"""
    from app.agents.orchestrator import OrchestratorAgent

    return await OrchestratorAgent().plan_route(
        content=content,
        file_url=file_urls[0] if file_urls else None,
        file_urls=file_urls,
        requested_agent=agent_type or "general",
    )


def _format_orchestrator_note(decision: dict) -> str:
    agent_name_map = {
        "resume": "简历评估 Agent",
        "recording": "录音分析 Agent",
        "question": "面试题生成 Agent",
        "general": "通用面试助手",
        "multi": "多 Agent 编排",
    }
    tasks = decision.get("tasks") if isinstance(decision.get("tasks"), list) else []
    agent_type = decision.get("agent_type", "general")
    if tasks:
        task_lines = []
        for index, task in enumerate(tasks, 1):
            task_lines.append(
                f"{index}. {agent_name_map.get(task.get('agent_type'), task.get('agent_type'))}"
                f"：{task.get('reason', '根据子任务需要调用。')}"
            )
        return (
            "## 主 Agent 调度\n"
            f"- **任务模式**：{'多 Agent 编排' if len(tasks) > 1 else '单 Agent 调用'}\n"
            f"- **调用模块**：{' → '.join(agent_name_map.get(t.get('agent_type'), str(t.get('agent_type'))) for t in tasks)}\n"
            f"- **判断依据**：{decision.get('reason', '根据用户当前请求进行调度。')}\n\n"
            "**子任务计划：**\n"
            + "\n".join(task_lines)
            + "\n\n---\n\n"
        )
    return (
        "## 主 Agent 调度\n"
        f"- **识别意图**：{decision.get('intent', 'interview_help')}\n"
        f"- **调用模块**：{agent_name_map.get(agent_type, agent_type)}\n"
        f"- **判断依据**：{decision.get('reason', '根据用户当前请求进行调度。')}\n\n"
        "---\n\n"
    )


async def _generate_sse_response(
    conversation_id: int,
    user_message: str,
    user_id: int,
    db: Session,
) -> str:
    """
    SSE 流式生成器
    调用 DeepSeek LLM 进行流式回复，支持思考模式和多轮对话上下文
    自动注入用户个人信息到 system prompt（如果用户允许）
    """
    from app.services.llm_service import LLMService

    llm = LLMService()

    # 构建多轮对话上下文：从数据库加载历史消息
    history_messages = (
        db.query(Message)
        .filter(Message.conversation_id == conversation_id)
        .order_by(Message.created_at.asc())
        .all()
    )

    # 将历史消息转换为 LLM messages 格式（只取 role + content）
    messages = []
    for msg in history_messages:
        # 当前用户消息已经在 history 中了（之前保存过），直接用
        messages.append({"role": msg.role, "content": msg.content})

    # 构建带用户个人上下文的 system prompt
    user_context = await _build_user_context(db, user_id)
    custom_prompt = SYSTEM_PROMPT + user_context if user_context else SYSTEM_PROMPT
    custom_prompt += (
        "\n\n当问题涉及具体岗位、公司、行业动态、公开答案、非通用知识或用户要求参考资料时，"
        "你需要结合联网搜索结果作答；引用网络资料时必须在回答末尾列出“参考来源”，包含标题和完整 URL。"
    )

    web_results = await _web_search(user_message, max_results=5)
    if web_results:
        refs = "\n".join(
            f"{i}. {item['title']}\nURL: {item['url']}\n摘要: {item['content']}"
            for i, item in enumerate(web_results, 1)
        )
        messages.append({
            "role": "system",
            "content": (
                "以下是联网搜索到的参考资料。请结合这些资料回答，并在回答末尾列出参考来源 URL：\n"
                f"{refs}"
            ),
        })
        yield json.dumps({
            "type": "status",
            "agent_name": "联网搜索",
            "message": f"已检索到 {len(web_results)} 条公开参考资料",
            "progress": 20,
        }, ensure_ascii=False)

    try:
        async for chunk in llm.chat_stream(
            messages=messages,
            system_prompt=custom_prompt,
        ):
            if chunk["type"] == "thinking":
                # 思考过程通过 thinking 事件发送
                yield json.dumps({
                    "type": "thinking",
                    "content": chunk["content"],
                    "done": False,
                }, ensure_ascii=False)
            elif chunk["type"] == "content":
                # 正式回复通过 message 事件发送
                async for text_chunk in _stream_content_chars(chunk["content"]):
                    yield text_chunk

    except Exception as e:
        logger.error(f"LLM 调用失败: {e}", exc_info=True)
        async for text_chunk in _stream_content_chars(f"[LLM 调用失败: {str(e)}]"):
            yield text_chunk

    # 发送完成信号
    yield json.dumps({
        "type": "content",
        "content": "",
        "done": True,
    }, ensure_ascii=False)


@router.post(
    "/conversations",
    response_model=APIResponse[ConversationResponse],
    summary="创建新会话",
)
async def create_conversation(
    request: ConversationCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """创建新的对话会话"""
    conversation = Conversation(
        user_id=current_user.id,
        title=request.title,
        agent_type=request.agent_type,
    )
    db.add(conversation)
    db.commit()
    db.refresh(conversation)

    return APIResponse(
        data=ConversationResponse.model_validate(conversation),
        message="会话创建成功",
    )


@router.get(
    "/conversations",
    response_model=APIResponse[ConversationListResponse],
    summary="获取会话列表",
)
async def list_conversations(
    page: int = Query(default=1, ge=1, description="页码"),
    page_size: int = Query(default=20, ge=1, le=100, description="每页数量"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取当前用户的会话列表"""
    from sqlalchemy import select, func as sa_func
    from app.models.message import Message

    # 子查询：每个会话的消息数量
    msg_count_subq = (
        select(
            Message.conversation_id,
            sa_func.count(Message.id).label("msg_count"),
        )
        .group_by(Message.conversation_id)
        .subquery()
    )

    # 子查询：每个会话最后一条消息的 ID（MySQL 兼容写法）
    last_msg_id_subq = (
        select(
            sa_func.max(Message.id).label("max_id"),
        )
        .group_by(Message.conversation_id)
        .subquery()
    )

    # 子查询：最后一条消息的内容
    last_msg_subq = (
        select(Message.id, Message.conversation_id, Message.content)
        .join(last_msg_id_subq, Message.id == last_msg_id_subq.c.max_id)
        .subquery()
    )

    # 查询总数
    total = (
        db.query(Conversation)
        .filter(Conversation.user_id == current_user.id)
        .count()
    )

    # 分页查询，附带 last_message 和 message_count
    rows = (
        db.query(
            Conversation,
            last_msg_subq.c.content.label("last_message"),
            sa_func.coalesce(msg_count_subq.c.msg_count, 0).label("message_count"),
        )
        .outerjoin(last_msg_subq, Conversation.id == last_msg_subq.c.conversation_id)
        .outerjoin(msg_count_subq, Conversation.id == msg_count_subq.c.conversation_id)
        .filter(Conversation.user_id == current_user.id)
        .order_by(Conversation.updated_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )

    result = []
    for conv, last_msg, msg_count in rows:
        conv_dict = ConversationResponse.model_validate(conv).model_dump()
        conv_dict["last_message"] = last_msg[:100] if last_msg else None
        conv_dict["message_count"] = msg_count
        result.append(ConversationResponse(**conv_dict))

    return APIResponse(
        data=ConversationListResponse(conversations=result, total=total),
    )


@router.get(
    "/conversations/{conversation_id}",
    response_model=APIResponse[ConversationDetailResponse],
    summary="获取会话详情",
)
async def get_conversation(
    conversation_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取指定会话的详细信息，包含消息列表"""
    conversation = (
        db.query(Conversation)
        .filter(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .first()
    )

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="会话不存在",
        )

    return APIResponse(
        data=ConversationDetailResponse.model_validate(conversation),
    )


@router.delete(
    "/conversations/{conversation_id}",
    response_model=APIResponse,
    summary="删除会话",
)
async def delete_conversation(
    conversation_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """删除指定会话及其所有消息"""
    conversation = (
        db.query(Conversation)
        .filter(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .first()
    )

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="会话不存在",
        )

    db.delete(conversation)
    db.commit()

    return APIResponse(message="会话已删除")


@router.get(
    "/conversations/{conversation_id}/messages",
    response_model=APIResponse[list[MessageResponse]],
    summary="获取消息历史",
)
async def get_messages(
    conversation_id: int,
    limit: int = Query(default=50, ge=1, le=200, description="返回消息数量"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取指定会话的消息历史记录"""
    # 验证会话归属
    conversation = (
        db.query(Conversation)
        .filter(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .first()
    )

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="会话不存在",
        )

    messages = (
        db.query(Message)
        .filter(Message.conversation_id == conversation_id)
        .order_by(Message.created_at.desc())
        .limit(limit)
        .all()
    )

    return APIResponse(
        data=[MessageResponse.model_validate(m) for m in reversed(messages)],
    )


@router.get(
    "/conversations/{conversation_id}/stream",
    summary="SSE 流式对话（GET，供 EventSource 调用）",
)
async def stream_chat(
    conversation_id: int,
    content: str = Query(..., description="消息内容"),
    message_type: str = Query(default="text", description="消息类型"),
    agent_type: str = Query(default="general", description="目标Agent类型：general/resume/recording/question"),
    file_url: str | None = Query(default=None, description="上传文件的URL（简历/录音）"),
    file_urls: str | None = Query(default=None, description="多个上传文件URL，逗号分隔"),
    file_names: str | None = Query(default=None, description="多个上传文件原始名称，逗号分隔"),
    token: str | None = Query(default=None, alias="token", description="JWT Token（EventSource 无法设置 Header，通过 query 传递）"),
    db: Session = Depends(get_db),
):
    """向指定会话发送消息，通过 SSE 流式返回 AI 回复"""
    # EventSource 不支持自定义 Header，需要从 query 参数获取 token
    if not token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="缺少认证令牌",
        )
    from app.api.auth import decode_token
    payload = decode_token(token)
    if payload.get("type") != "access":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="令牌类型错误",
        )
    user_id = payload.get("sub")
    from app.models.user import User
    current_user = db.query(User).filter(User.id == int(user_id)).first()
    if not current_user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="用户不存在",
        )

    uploaded_file_urls = _parse_file_urls(file_url, file_urls)
    uploaded_file_names = _parse_file_names(file_names, uploaded_file_urls)
    logger.info(
        "SSE 请求: conversation_id=%s, agent_type=%s, file_urls=%s, content=%s",
        conversation_id,
        agent_type,
        uploaded_file_urls,
        content[:50] if content else "(empty)",
    )
    agent_decision = await _resolve_agent_decision(agent_type, content, uploaded_file_urls)
    resolved_agent_type = agent_decision.get("agent_type", "general")

    # 验证会话归属
    conversation = (
        db.query(Conversation)
        .filter(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .first()
    )

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="会话不存在",
        )

    existing_message_count = db.query(Message).filter(Message.conversation_id == conversation_id).count()
    if existing_message_count == 0 and (not conversation.title or conversation.title == "新对话"):
        conversation.title = _derive_conversation_title(content)

    # 保存用户消息
    user_msg = Message(
        conversation_id=conversation_id,
        role="user",
        content=content,
        message_type="file" if uploaded_file_urls else (message_type or "text"),
        metadata_=({
            "file_url": uploaded_file_urls[0],
            "file_urls": uploaded_file_urls,
            "file_name": "、".join(uploaded_file_names),
            "file_names": uploaded_file_names,
        } if uploaded_file_urls else None),
    )
    db.add(user_msg)
    db.commit()

    # 构建 SSE 事件流
    async def event_generator():
        full_response = ""
        start_time = time.perf_counter()
        try:
            orchestration_note = _format_orchestrator_note(agent_decision)
            async for note_chunk in _stream_content_chars(orchestration_note, agent_name="主Agent"):
                yield {"event": "message", "data": note_chunk}
                full_response += json.loads(note_chunk).get("content", "")

            yield {"event": "message", "data": json.dumps({
                "type": "status",
                "agent_name": "主Agent",
                "message": agent_decision.get("reason", "主 Agent 已完成意图分析"),
                "progress": 10,
            }, ensure_ascii=False)}

            async def forward_agent_stream(stream):
                nonlocal full_response
                async for chunk_json in stream:
                    chunk_data = json.loads(chunk_json)
                    if chunk_data.get("done"):
                        continue
                    yield {"event": "message", "data": chunk_json}
                    if chunk_data.get("type") == "content":
                        full_response += chunk_data.get("content", "")

            tasks = agent_decision.get("tasks") if isinstance(agent_decision.get("tasks"), list) else []
            if not tasks:
                tasks = [{
                    "agent_type": resolved_agent_type,
                    "intent": agent_decision.get("intent", "interview_help"),
                    "reason": agent_decision.get("reason", "主 Agent 选择单 Agent 处理。"),
                    "file_urls": uploaded_file_urls[:1],
                }]

            for index, task in enumerate(tasks, 1):
                task_agent = task.get("agent_type", "general")
                task_file_url = _select_file_for_agent(task, uploaded_file_urls)
                task_title = {
                    "resume": "简历评估 Agent",
                    "recording": "录音分析 Agent",
                    "question": "面试题生成 Agent",
                    "general": "通用面试助手",
                }.get(task_agent, str(task_agent))

                if len(tasks) > 1:
                    section = f"\n\n## 子任务 {index}：{task_title}\n\n{task.get('reason', '')}\n\n"
                    async for section_chunk in _stream_content_chars(section, agent_name="主Agent"):
                        yield {"event": "message", "data": section_chunk}
                        full_response += json.loads(section_chunk).get("content", "")

                yield {"event": "message", "data": json.dumps({
                    "type": "status",
                    "agent_name": task_title,
                    "message": task.get("reason", "子 Agent 正在处理..."),
                    "progress": min(95, 15 + index * 20),
                }, ensure_ascii=False)}

                if task_agent == "resume" and task_file_url:
                    async for event in forward_agent_stream(_generate_resume_sse(
                        conversation_id, task_file_url, content, current_user.id, db
                    )):
                        yield event
                elif task_agent == "recording" and task_file_url:
                    async for event in forward_agent_stream(_generate_recording_sse(
                        conversation_id, task_file_url, content, current_user.id, db
                    )):
                        yield event
                elif task_agent == "question":
                    async for event in forward_agent_stream(_generate_question_sse(
                        conversation_id, content, current_user.id, db, task_file_url
                    )):
                        yield event
                elif task_agent == "general":
                    if uploaded_file_urls:
                        warning = (
                            "\n\n> 主 Agent 说明：通用面试助手没有原始文件解析权限。"
                            "涉及文件的任务会优先交给简历、录音或题目 Agent；当前仅处理纯文本部分。\n\n"
                        )
                        async for warning_chunk in _stream_content_chars(warning, agent_name="主Agent"):
                            yield {"event": "message", "data": warning_chunk}
                            full_response += json.loads(warning_chunk).get("content", "")
                    async for event in forward_agent_stream(_generate_sse_response(
                        conversation_id, content, current_user.id, db
                    )):
                        yield event
                else:
                    missing = f"\n\n[{task_title} 无法执行：没有匹配到可处理的上传文件]\n\n"
                    async for missing_chunk in _stream_content_chars(missing, agent_name="主Agent"):
                        yield {"event": "message", "data": missing_chunk}
                        full_response += json.loads(missing_chunk).get("content", "")

            # 保存 AI 回复到数据库（只保存正式回复，不含思考过程）
            from app.database import SessionLocal
            save_db = SessionLocal()
            try:
                ai_msg = Message(
                    conversation_id=conversation_id,
                    role="assistant",
                    content=full_response,
                    agent_name=resolved_agent_type,
                    message_type="text",
                )
                save_db.add(ai_msg)
                from app.models.agent_log import AgentLog

                save_db.add(AgentLog(
                    conversation_id=conversation_id,
                    agent_name=resolved_agent_type,
                    input_tokens=None,
                    output_tokens=None,
                    latency_ms=int((time.perf_counter() - start_time) * 1000),
                    status="success",
                    error_msg=None,
                ))
                save_db.commit()
            finally:
                save_db.close()

            yield {"event": "done", "data": "{}"}

        except Exception as e:
            logger.error(f"SSE 事件流异常: {e}", exc_info=True)
            try:
                from app.database import SessionLocal
                from app.models.agent_log import AgentLog

                log_db = SessionLocal()
                try:
                    log_db.add(AgentLog(
                        conversation_id=conversation_id,
                        agent_name=resolved_agent_type,
                        input_tokens=None,
                        output_tokens=None,
                        latency_ms=int((time.perf_counter() - start_time) * 1000),
                        status="failed",
                        error_msg=str(e),
                    ))
                    log_db.commit()
                finally:
                    log_db.close()
            except Exception:
                pass
            yield {
                "event": "error",
                "data": json.dumps({"error": str(e)}, ensure_ascii=False),
            }

    return EventSourceResponse(event_generator())


@router.get(
    "/conversations/{conversation_id}/artifacts",
    response_model=APIResponse[list[dict]],
    summary="获取会话产物列表",
)
async def list_artifacts(
    conversation_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conversation = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == current_user.id)
        .first()
    )
    if not conversation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="会话不存在")

    directory = _artifact_dir(current_user.id, conversation_id)
    items = []
    for path in sorted(directory.iterdir(), key=lambda item: item.stat().st_mtime, reverse=True):
        if not path.is_file():
            continue
        stat = path.stat()
        items.append({
            "name": path.name,
            "size": stat.st_size,
            "type": path.suffix.lower().lstrip(".") or "file",
            "created_at": stat.st_mtime,
            "preview_url": f"/api/v1/chat/conversations/{conversation_id}/artifacts/{path.name}/preview",
            "download_url": f"/api/v1/chat/conversations/{conversation_id}/artifacts/{path.name}/download",
        })
    return APIResponse(data=items)


@router.get(
    "/conversations/{conversation_id}/artifacts/{filename}/preview",
    summary="预览会话产物",
)
async def preview_artifact(
    conversation_id: int,
    filename: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conversation = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == current_user.id)
        .first()
    )
    if not conversation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="会话不存在")
    safe_name = Path(filename).name
    path = _artifact_dir(current_user.id, conversation_id) / safe_name
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="产物不存在")
    return PlainTextResponse(_preview_artifact_text(path))


@router.get(
    "/conversations/{conversation_id}/artifacts/{filename}/download",
    summary="下载会话产物",
)
async def download_artifact(
    conversation_id: int,
    filename: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conversation = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == current_user.id)
        .first()
    )
    if not conversation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="会话不存在")
    safe_name = Path(filename).name
    path = _artifact_dir(current_user.id, conversation_id) / safe_name
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="产物不存在")
    return FileResponse(path, filename=safe_name)


@router.post(
    "/upload",
    summary="上传文件（简历/录音）",
)
async def upload_file(
    file: "UploadFile",
    current_user: User = Depends(get_current_user),
):
    """上传文件到本地存储，返回文件URL"""
    import os
    import uuid
    from pathlib import Path
    from fastapi import UploadFile

    # 创建上传目录
    upload_dir = Path("uploads") / str(current_user.id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    # 生成唯一文件名
    ext = Path(file.filename).suffix if file.filename else ""
    unique_name = f"{uuid.uuid4().hex}{ext}"
    file_path = upload_dir / unique_name

    # 保存文件
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # 返回文件URL（相对路径，前端通过后端静态文件访问）
    file_url = f"/uploads/{current_user.id}/{unique_name}"

    return APIResponse(
        data={
            "file_url": file_url,
            "file_name": file.filename,
            "file_size": len(content),
            "file_type": ext.lstrip(".").lower(),
        },
        message="文件上传成功",
    )
